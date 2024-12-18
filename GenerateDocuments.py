from abc import ABC, abstractmethod
from docxtpl import DocxTemplate
from docx import Document
import openpyxl
from openpyxl.styles import Font
from Constants.FileNames import FileNames
from Constants.ExcelFields import ExcelFields

class IGenerateFile(ABC):
    @abstractmethod
    def generate_file(self, data: list | str) -> None:
        pass
    
class ICreator:
    @abstractmethod
    def create(self) -> IGenerateFile:
        pass


class GenerateRequestDocx(IGenerateFile):
    def generate_file(self, data: list | str) -> None:
        try:
            doc = DocxTemplate("Письмо.docx")
            context = {'name': data}
            doc.render(context)
            doc.save(FileNames.RequestFileName)
            print("Файл успешно сохранен.")
        except Exception as e:
            print("Ошибка при сохранении файла:", e)


class GenerateTechDevicesDocx(IGenerateFile):
    def generate_file(self, data: list | str) -> None:
        """Функция для заполнения технических средств с учётом количества"""
        doc = Document('example_passport.docx')

        # Доступ к таблице
        table = doc.tables[0]
        for i, item in enumerate(data):
            # Добавьте новую строку в таблицу
            row = table.add_row()

            row.cells[0].text = str(i + 1)  # Номер строки
            row.cells[1].text = item['name_sys']
            row.cells[2].text = item['model_sys']
            row.cells[3].text = item['id_sys']
            row.cells[4].text = f"{item['address']}, кб. № {item['office_num']}"

            # Добавить новую строку в конец таблицы
            last_row = table.add_row()
            # Объединить все ячейки в новой строке
            table.cell(last_row._index, 0).merge(table.cell(last_row._index, table._column_count - 1))
            # Добавить текст в объединенную ячейку
            last_row.cells[0].add_paragraph("Периферийное оборудование: клавиатура, \
                                            манипулятор <мышь>, монитор, принтер, WEB-камера, колонки.")

        doc.save(FileNames.TechPassportFileName)


class GenerateExcelARM(IGenerateFile):
    def generate_file(self, data: list | str) -> None:
        workbook = openpyxl.load_workbook("list2.xlsx")
        # Выбор активного листа
        sheet = workbook.active
        for index, value in enumerate(ExcelFields.ExcelARMFields):
            sheet[value].font = Font(name='Times New Roman', size=14)
            sheet[value] = data[index]

        workbook.save(FileNames.ExcelARMName)


class GenerateExcelConnectARM(IGenerateFile):
    def generate_file(self, data: list | str) -> None:
        workbook = openpyxl.load_workbook("list4.xlsx")
        # Выбор активного листа
        sheet = workbook.active
        for index, value in enumerate(ExcelFields.ExcelConnectARMFields):
            sheet[value].font = Font(name='Times New Roman', size=14)
            sheet[value] = data[index]

        workbook.save(FileNames.RequestToConnectARMFileName)


class GenerateExcelList1(IGenerateFile):
    def generate_file(self, data: list | str) -> None:
        workbook = openpyxl.load_workbook("list1.xlsx")
        # Выбор активного листа
        sheet = workbook.active
        for index, value in enumerate(ExcelFields.ExcelList1Fields):
            sheet[value].font = Font(name='Times New Roman', size=14)
            sheet[value] = data[index]

        workbook.save(f'final_list1.xlsx')


class RequestFileCreator(ICreator):
    def create(self) -> IGenerateFile:
        return GenerateRequestDocx()


class TechDevicesFileCreator(ICreator):
    def create(self) -> IGenerateFile:
        return GenerateTechDevicesDocx()


class ExcelARMCreator(ICreator):
    def create(self) -> IGenerateFile:
        return GenerateExcelARM()

class ExcelConnectARMCreator(ICreator):
    def create(self) -> IGenerateFile:
        return GenerateExcelConnectARM()

class ExcelList1Creator(ICreator):
    def create(self) -> IGenerateFile:
        return GenerateExcelList1()





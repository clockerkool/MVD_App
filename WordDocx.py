from docxtpl import DocxTemplate
from docx import Document

def get_tech(name_sys, model_sys, id_sys, address, office_num):
    "Функция для заполнения технических средств без учёта количества(больше не нужна)"
    print(name_sys, model_sys, id_sys, address, office_num)
    doc = DocxTemplate("example_passport.docx")
    context = { 'name_sys' : name_sys, 'model_sys': model_sys, 'id_sys': id_sys,
                'address': address, 'office_num': office_num
                }
    doc.render(context)
    doc.save("tech_passport.docx")



def get_tech2(data):
    """Функция для заполнения технических средств с учётом количества"""
    doc = Document('example_passport.docx')

    # Доступ к таблице
    table = doc.tables[0]
    for i, item in enumerate(data):
        # Добавьте новую строку в таблицу
        row = table.add_row()

        row.cells[0].text = str(i + 1)  # Номер строки
        row.cells[1].text = item['name_sys']
        row.cells[2].text = item['model_sys']
        row.cells[3].text = item['id_sys']
        row.cells[4].text = f"{item['address']}, кб. № {item['office_num']}"

        # Добавить новую строку в конец таблицы
        last_row = table.add_row()
        # Объединить все ячейки в новой строке
        table.cell(last_row._index, 0).merge(table.cell(last_row._index, table._column_count - 1))
        # Добавить текст в объединенную ячейку
        last_row.cells[0].add_paragraph(
            "Периферийное оборудование: клавиатура, манипулятор <мышь>, монитор, принтер, WEB-камера, колонки.")

    doc.save('filled_passport.docx')


def get_request(name):
    try:
        doc = DocxTemplate("Письмо.docx")
        context = { 'name' : name}
        doc.render(context)
        print("ALL RIGHT")
        doc.save("Заявление.docx")
        print("Файл успешно сохранен.")
    except Exception as e:
        print("Ошибка при сохранении файла:", e)






